import io
import re
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from reportlab.lib.units import inch
from django.core.files.base import ContentFile

class DocumentGenerator:
    @staticmethod
    def clean_text(text):
        """
        Cleans text to prevent encoding issues in PDF/DOCX.
        Replaces smart quotes/dashes and removes unsupported Unicode.
        """
        if not text:
            return ""
            
        # Replace common 'smart' characters that cause issues with standard fonts
        replacements = {
            '\u2018': "'", '\u2019': "'",  # Smart single quotes
            '\u201c': '"', '\u201d': '"',  # Smart double quotes
            '\u2013': '-', '\u2014': '-',  # En/Em dashes
            '\u2022': '*',                  # Bullets
            '\u2026': '...',                # Ellipsis
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
            
        # Remove other non-ASCII characters that might cause issues in PDF Helvetica
        text = "".join(i for i in text if ord(i) < 128)
        return text

    @staticmethod
    def generate_docx(text):
        """Generates a professional .docx file from text."""
        text = DocumentGenerator.clean_text(text)
        doc = Document()
        
        # Set margins (standard 1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add content with proper paragraph handling
        paragraphs = text.split('\n')
        for p_text in paragraphs:
            p_text = p_text.strip()
            if p_text:
                p = doc.add_paragraph(p_text)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
            else:
                doc.add_paragraph() # Spacer

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_pdf(text):
        """Generates a professional .pdf file using Platypus for wrapping/layout."""
        text = DocumentGenerator.clean_text(text)
        buffer = io.BytesIO()
        
        # Create the document template
        doc = SimpleDocTemplate(
            buffer,
            pagesize=LETTER,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        styles = getSampleStyleSheet()
        
        # Custom style for the body text
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=10
        )
        
        # Build the story (list of elements)
        story = []
        
        # Split text into segments based on double newlines for paragraphs
        # or single newlines for headers/lists
        segments = text.split('\n')
        for segment in segments:
            segment = segment.strip()
            if segment:
                # Escape HTML-like characters because Paragraph treats text as pseudo-HTML
                segment = segment.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(segment, body_style))
            else:
                story.append(Spacer(1, 12))
        
        # Generate the PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def get_document_content(text, format_type):
        if format_type.lower() == 'docx':
            return DocumentGenerator.generate_docx(text)
        else:
            return DocumentGenerator.generate_pdf(text)
